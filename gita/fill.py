from docxtpl import DocxTemplate
from datetime import date
import docx
from docx.shared import Pt
from pyBSDate import convert_AD_to_BS
import os
import regex

cwd = os.getcwd()

# pip install pyBSDate
# pip install docxtpl
# pip install Pillow
# pip install regex
# pip install django

def fill_my_form(main_dir,
                photo_ = 'photo.png',
                sign_ = 'sign.png',
                regis = '123456-12', 
                usr_year = int('5'),
                usr_sem = int('2'),
                eng_name = 'Ram Prasad',
                nep_name = 'राम प्रसाद',
                # first_n = 'John',
                # middle_n = '',
                # last_n = 'Snowy',
                # nepali_fname = 'सेदुम्बा',
                # nepali_mname = 'सर्सापी',
                # nepali_lname = 'अर्याल',
                year_e = '1998',
                month_e = '1',
                day_e = '1',
                usr_gender = 'female', # options: 'male', 'female', 'any_string_really'
                ):

    if len(eng_name.strip().split()) == 2:
        first_n, last_n = eng_name.strip().split()
        middle_n = ''
    elif len(eng_name.strip().split()) == 3:
        first_n, middle_n, last_n = eng_name.strip().split()
    else:
        first_n, middle_n, last_n = ['wrong', '', 'name']

    if len(nep_name.strip().split()) == 2:
        nepali_fname, nepali_lname = nep_name.strip().split()
        nepali_mname = ''
    elif len(nep_name.strip().split()) == 3:
        nepali_fname, nepali_mname, nepali_lname = nep_name.strip().split()
    else:
        nepali_fname, nepali_mname, nepali_lname = ['गलत', '', 'नाम']

    






    ph_locate = os.path.join(main_dir, 'fill_files', photo_)
    sgn_locate = os.path.join(main_dir, 'fill_files', sign_)

    # All about time
    today = date.today()
    today = today.strftime("%B %d, %Y")
    this_year = today.split(',')[1][1:]
    year_one = str(int(this_year) - 1)[-2:]
    year_two = this_year[-2:]


    #Loading the template via docxtemplate modules
    template_path = os.path.join(main_dir, 'fill_files', "template.docx")
    doc = DocxTemplate(template_path)


    # !!!!!!!   year section
    if usr_year == 1:
        usr_year_s = 'st'
    elif usr_year == 2:
        usr_year_s = 'nd'
    elif usr_year == 3:
        usr_year_s = 'rd'
    else:
        usr_year_s = 'th'

    # !!!!!!!!    sem section
    if usr_sem == 1:
        usr_sem_s = 'st'
    elif usr_sem == 2:
        usr_sem_s = 'nd'

    batcher = '20' + regis[-2:]

    # All about photos
    doc.replace_pic("dummy_photo.png", ph_locate)
    doc.replace_pic("dummy_signature.png", sgn_locate)

    # subjects section: It will depend on what the usr_year and usr_sem
    if usr_year == 1 and usr_sem == 1:
        subject = {1: ['English I ( General)', 'LENG', '101'],
                   2: ['Jurisprudence', 'LJUR', '111'],
                   3: ['Constitutional Law I', 'LCNS', '121'],
                   4: ['Contract Law', 'LCON', '161'],
                   5: ['Financial Accounting', 'ACC', '205'],
                   6: ['Business Management', 'GEM', '231']}
    if usr_year == 1 and usr_sem == 2:
        subject = {1: ['Microeconomics', 'ECO', '201'],
                   2: ['Nepali (General)', 'LNEP', '103'],
                   3: ['Competition Law', 'LCMP', '162'],
                   4: ['English II (Legal English)', 'LENG', '102'],
                   5: ['Quantitative Techniques', 'MAS', '103'],
                   6: ['Constitutional Law II', 'LCNS', '122']}
    if usr_year == 2 and usr_sem == 1:
        subject = {1: ['Managerial Communication', 'GEM', '201'],
                   2: ['Nepali II (Legal Language)', 'LNEP', '204'],
                   3: ['Management Accounting', 'ACC', '210'],
                   4: ['Human Rights Law and Practice', 'LHRT', '223'],
                   5: ['Torts & Consumer Protection Laws', 'LCPT', '231'],
                   6: ['Family Law', 'LFAM', '232']}
    if usr_year == 2 and usr_sem == 2:
        subject = {1: ['Macroeconomics', 'ECO', '210'],
                   2: ['Legal Reasoning Skill & Logic', 'LRSK', '233'],
                   3: ['Law of Crimes I (Penal Code)', 'LCRM', '241'],
                   4: ['Organizational Behavior', 'HRM', '320'],
                   5: ['Cyber Law', 'LCYB', '234'],
                   6: ['Philosophy of Life & Life Style', 'LPHI', '205']}
    if usr_year == 3 and usr_sem == 1:
        subject = {1: ['Human Resource Management', 'HRM', '201'],
                   2: ['Financial Management', 'FIN', '301'],
                   3: ['Property Law', 'LPRT', '335'],
                   4: ['Law of Crimes II (Criminal Procedure Code I)', 'LCRM', '341'],
                   5: ['Company Law I', 'LCOM', '351'],
                   6: ['Administration Law', 'LADM', '336']}
    if usr_year == 3 and usr_sem == 2:
        subject = {1: ['Environmental law', 'LENV', '371'],
                   2: ['Intellectual Property Law', 'LINP', '368'],
                   3: ['Marketing Management', 'MKT', '310'],
                   4: ['Company Law II', 'LCOM', '352'],
                   5: ['Law of Crimes II', 'LCRM', '342'],
                   6: ['International Business', 'GEM', '470']}
    if usr_year == 4 and usr_sem == 1:
        subject = {1: ['Project Management', 'GEM', '332'],
                   2: ['Banking Law', 'LBNK', '453'],
                   3: ['Public International Law', 'LINT', '472'],
                   4: ['Civil Procedure & Limitation Law I', 'LPCD', '437'],
                   5: ['Financial Institutions & Markets', 'FIN', '330'],
                   6: ['Merger & Acquisitions Law', 'LMRG', '454']}
    if usr_year == 4 and usr_sem == 2:
        subject = {1: ['Civil Procedure and Limitation Law II', 'LPCD', '439'],
                   2: ['Taxation Law', 'LTAX', '464'],
                   3: ['Corporate Governance and Business Ethics', 'LCGB', '481'],
                   4: ['Law of Evidence', 'LEVD', '438'],
                   5: ['Strategic Management', 'GEM', '490'],
                   6: ['Insurance Law', 'LINS', '463']}
    if usr_year == 5 and usr_sem == 1:
        subject = {1: ['Entrepreneurship and New Business Formation', 'GEM', '310'],
                   2: ['Water and Energy Law', 'LWTR', '573'],
                   3: ['Labour and Industrial Law', 'LIND', '565'],
                   4: ['Investment Law', 'LINV', '566'],
                   5: ['Legal Research Analysis and Writing', 'LDBS', '506'],
                   6: ['Trade Law', 'LTRD', '567']}
    if usr_year == 5 and usr_sem == 2:
        subject = {1: ['Drafting, Pleading and Conveyance', 'LDPC', '591'],
                   2: ['Alternative Dispute Resolution', 'LADR', '592'],
                   3: ['Professional Ethics and Professional Accounting System', 'LPEA', '593'],
                   4: ['Moot Court Exercise and Internship', 'LMTC', '594'],
                   5: ['Court Planning and Management', 'LCPM', '595'],
                   6: ['\n', '', '']}

    # fix for the 10th sem
    if usr_year == 5 and usr_sem == 2:
        fix = ''
    else:
        fix = '3'

    context = {
        # the subjects
        'one': f'{subject[1][0]}',
        'two': f'{subject[2][0]}',
        'three': f'{subject[3][0]}',
        'four': f'{subject[4][0]}',
        'five': f'{subject[5][0]}',
        'six': f'{subject[6][0]}',
        'one_c1': f'{subject[1][1]}',
        'two_c1': f'{subject[2][1]}',
        'three_c1': f'{subject[3][1]}',
        'four_c1': f'{subject[4][1]}',
        'five_c1': f'{subject[5][1]}',
        'six_c1': f'{subject[6][1]}',
        'one_c2': f'{subject[1][2]}',
        'two_c2': f'{subject[2][2]}',
        'three_c2': f'{subject[3][2]}',
        'four_c2': f'{subject[4][2]}',
        'five_c2': f'{subject[5][2]}',
        'six_c2': f'{subject[6][2]}',

        # reg, name, batch, year, sem, today's date
        'reg' : f'{regis}',
        'f_name': f'{first_n.capitalize()}',
        'm_name': f'{middle_n.capitalize()}',
        'l_name': f'{last_n.capitalize()}',
        'batch': f'{batcher}',
        'year_s' : f'{usr_year}',
        'year_super': f'{usr_year_s}',
        'sem_s': f'{usr_sem}',
        'sem_super': f'{usr_sem_s}',
        'today_date': f'{today}',
        'y_1': f'{year_one}',
        'y_2': f'{year_two}',
        'fix_10': f'{fix}'}

    # Saving the channges
    docx_file_name = os.path.join(main_dir, 'fill_files', f'{regis}.docx')

    doc.render(context)
    doc.save(docx_file_name)


    # Using the next module
    d = docx.Document(docx_file_name)

    

    the_name = last_n.upper() + ' ' + first_n.upper() + ' ' + middle_n.upper()
    the_name = the_name.strip()

    def fix_font(run):
        font = run.font
        select_font = 'Times New Roman'
        font.name = select_font
        font.size = Pt(12)

    # ????? here, if len(the_name) is above certain number, maybe throw an error
    for index, ch in enumerate(the_name):
        d.tables[6].cell(0, index).text = ch
        try:
            c = d.tables[6].cell(0, index).paragraphs[0].run[0]
            fix_font(c)
        except:
            pass

    # !!!!!!!   All about the Name in Nepali (pg 2)
    nepali_fname = nepali_fname.split(' ')
    nepali_mname = nepali_mname.split(' ')
    nepali_lname = nepali_lname.split(' ')

    f = the_name.split(first_n.upper())

    # determination of space in name having middle name
    if len(f[1]) > 1:
        # Space after last name
        first_space = 0
        second_space = 0
        for i in f[0]:
            if i == ' ':
                first_space += 1
        # Space after first name
        for i in f[1]:
            if i == ' ':
                second_space += 1

    # determination of space in name not having middle name
    elif len(f[1]) < 2:
        # Space after last name
        first_space = 0
        second_space = 0
        for i in f[0]:
            if i == ' ':
                first_space += 1

    # determination of space in name which is somehow messed up
    else:
        if nepali_mname == '':
            first_space = 6
            second_space = 3
        else:
            first_space = 5
            second_space = 3

    first_space = len(last_n) + first_space - len(nepali_lname)
    second_space = len(first_n) + second_space - len(nepali_fname)

    def int_to_list(number):
        my_list = []
        for i in range(number):
            my_list.append('')
        return my_list


    first_space = int_to_list(first_space)
    second_space = int_to_list(second_space)

    # Method 1: uncomment the below one line. Note: The rest of this method was left uncommented since it wouldn't affect method 2 and reviving method 1 would be easy.
    # nepali = nepali_lname + first_space + ['',] + nepali_fname + second_space + ['',] + nepali_mname

    # Method 2: in this method, the user doesn't have to separate nepali characters. To use method 1, comment below upto nepali = ...
    nepali_fname = ' '.join(nepali_fname)
    nepali_mname = ' '.join(nepali_mname)
    nepali_lname = ' '.join(nepali_lname)
    names = [regex.findall(r'\X', i) for i in [str(nepali_lname), str(nepali_fname), str(nepali_mname)]]
    nepali_lname, nepali_fname, nepali_mname = [names[0], names[1], names[2]] 
    nepali_lname.append('')
    nepali_fname.append('')
    nepali = nepali_lname + nepali_fname + nepali_mname


    for index, ch in enumerate(nepali):
        d.tables[6].cell(1, index).text = ch
        try:
            c = d.tables[6].cell(1, index).paragraphs[0].run[0]
            fix_font(c)
        except:
            pass


    # messy stuff part
    to_convert_year = int(year_e)
    to_convert_month = int(month_e)
    to_convert_day = int(day_e)

    if len(month_e) == 1:
        month_e = '0' + month_e

    if len(day_e) == 1:
        day_e = '0' + day_e

    DOB_E = year_e + month_e + day_e
    for index, ch in enumerate(DOB_E):
        index += 9
        d.tables[7].cell(1, index).text = ch
        try:
            c = d.tables[7].cell(1, index).paragraphs[0].run[0]
            fix_font(c)
        except:
            pass

    bs_date = convert_AD_to_BS(to_convert_year, to_convert_month, to_convert_day)
    year_n, month_n, day_n = str(bs_date[0]), str(bs_date[1]), str(bs_date[2])

    if len(month_n) == 1:
        month_n = '0' + month_n

    if len(day_n) == 1:
        day_n = '0' + day_n

    DOB_N = year_n + month_n + day_n
    for index, ch in enumerate(DOB_N):
        d.tables[7].cell(1, index).text = ch
        try:
            c = d.tables[7].cell(1, index).paragraphs[0].run[0]
            fix_font(c)
        except:
            pass


    # Block letter registration
    for index, ch in enumerate(regis):
        d.tables[4].cell(0, index).text = ch
        try:
            c = d.tables[7].cell(1, index).paragraphs[0].run[0]
            fix_font(c)
        except:
            pass


    # !!! Give a scroll of the options of 'tick on Male', 'tick on Female'. Maybe add an option of 'don't tick'.

    p = d.paragraphs[30]
    if usr_gender in ['m', 'male', 'tick male']:
        p.add_run(':  Male✔       Female')

    elif usr_gender in ['f', 'female', 'tick female']:
        p.add_run(':  Male        Female✔')

    else:
        p.add_run(':  Male        Female')

    c = d.paragraphs[30].runs[1]
    fix_font(c)

    d.save(docx_file_name)
    print(docx_file_name)

    # file_types = [("PNG file", "*.png"), ("EMG file", "*.emg"), ("JPG file", "*.jpg"), ("JPG file2", "*.jpeg"),
    # ("WMF file", "*.wmf"), ("TIFF file", "*.tif"), ("TIFF file2", "*.tiff")]

if __name__ == "__main__":
    BASE_DIR = '/Users/birajaryal/Desktop/quizy/mysite-project'
    fill_my_form(BASE_DIR, 'empty_photo.png', eng_name='Biraj Kumar Aryal', nep_name='विराज कुमार अर्याल')
